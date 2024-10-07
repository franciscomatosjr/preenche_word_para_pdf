from docx import Document
import comtypes.client
from PyPDF2 import PdfReader, PdfWriter


# Passo 1: Preencher o documento Word com as tags
def preencher_tags_no_documento(modelo_path, dados, output_path):
    doc = Document(modelo_path)
    for paragrafo in doc.paragraphs:
        for tag, valor in dados.items():
            if f"{{{{{tag}}}}}" in paragrafo.text:
                paragrafo.text = paragrafo.text.replace(f"{{{{{tag}}}}}", valor)

    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for tag, valor in dados.items():
                    if f"{{{{{tag}}}}}" in celula.text:
                        celula.text = celula.text.replace(f"{{{{{tag}}}}}", valor)

    doc.save(output_path)


# Passo 2: Converter o documento Word preenchido em PDF
def converter_word_para_pdf(word_path, pdf_path):
    word = comtypes.client.CreateObject('Word.Application')
    word.Visible = False
    doc = word.Documents.Open(word_path)
    doc.SaveAs(pdf_path, FileFormat=17)  # 17 é o formato para PDF
    doc.Close()
    word.Quit()


# Passo 3: Substituir a primeira página de um PDF por uma nova
def substituir_primeira_pagina_pdf(pdf_original_path, nova_pagina_pdf, output_pdf_path):
    original_pdf = PdfReader(pdf_original_path)
    nova_pagina = PdfReader(nova_pagina_pdf)

    writer = PdfWriter()

    writer.add_page(nova_pagina.pages[0])
    for page_num in range(1, len(original_pdf.pages)):
        writer.add_page(original_pdf.pages[page_num])

    with open(output_pdf_path, 'wb') as output_file:
        writer.write(output_file)


# Exemplo de uso
if __name__ == "__main__":
    # Arquivos
    modelo_word = r"C:\Users\Francisco\Downloads\PAG01 - FICHA DE CADASTRO BENEFICIARIO - COLETIVO POR ADESÃO - VOCE TOTAL.docx"  # Caminho para o arquivo modelo
    word_preenchido = r'C:\Users\Francisco\Downloads\documento_preenchido1.docx'  # Caminho de saída para o novo documento

    # modelo_word = 'modelo_com_tags.docx'
    # word_preenchido = 'documento_preenchido.docx'
    pdf_preenchido = r'C:\Users\Francisco\Downloads\primeira_pagina.pdf'
    pdf_original = r'C:\Users\Francisco\Downloads\PROPOSTA COMPLETA - VOCÊ TOTAL ADESÃO.pdf'
    pdf_final = r'C:\Users\Francisco\Downloads\novo_pdf_com_pagina_substituida.pdf'

    # Dados a serem inseridos
    dados = {
        "NOME": "Maria Oliveira",
        "ENDERECO": "Avenida Paulista, 1000",
        "DATA": "7 de Novembro de 2024"
    }

    # Passo 1: Preencher o documento Word com as tags
    preencher_tags_no_documento(modelo_word, dados, word_preenchido)

    # Passo 2: Converter o documento Word preenchido em PDF
    converter_word_para_pdf(word_preenchido, pdf_preenchido)

    # Passo 3: Substituir a primeira página do PDF original
    substituir_primeira_pagina_pdf(pdf_original, pdf_preenchido, pdf_final)
